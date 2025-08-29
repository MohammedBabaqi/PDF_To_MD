import os
import io
import time
import base64
from concurrent.futures import ThreadPoolExecutor

import streamlit as st

try:
    # python-docx is optional; install it if you want DOCX output
    from docx import Document as DocxDocument
    from docx.shared import Inches
    HAS_PYTHON_DOCX = True
except Exception:
    HAS_PYTHON_DOCX = False

from mistralai import Mistral, DocumentURLChunk

# -----------------------------
# Helpers
# -----------------------------
def get_api_key() -> str:
    """Fetch API key from Streamlit Secrets (preferred) or environment var."""
    # 1) Streamlit Cloud / local .streamlit/secrets.toml
    key = st.secrets.get("MISTRAL_API_KEY", None) if hasattr(st, "secrets") else None
    # 2) Fallback to environment variable (for local dev)
    if not key:
        key = os.getenv("MISTRAL_API_KEY")
    return key 


def pdf_bytes_to_data_url(pdf_bytes: bytes) -> str:
    b64 = base64.b64encode(pdf_bytes).decode("utf-8")
    return f"data:application/pdf;base64,{b64}"


def run_ocr(client: Mistral, pdf_bytes: bytes, include_images: bool = True):
    data_url = pdf_bytes_to_data_url(pdf_bytes)
    return client.ocr.process(
        model="mistral-ocr-latest",
        document=DocumentURLChunk(document_url=data_url),
        include_image_base64=include_images,
    )


def response_to_markdown(resp) -> str:
    """Join page markdown; append extracted images as embedded data URIs per page."""
    md_parts = []
    for p in resp.pages:
        # Page header
        md_parts.append(f"\n\n---\n\n### Page {p.index + 1}\n\n")
        # Page text
        if getattr(p, "markdown", None):
            md_parts.append(p.markdown)
        # Images (if included)
        images = getattr(p, "images", []) or []
        if images:
            md_parts.append("\n\n#### Extracted images\n")
            for i, im in enumerate(images, start=1):
                b64 = getattr(im, "image_base64", None)
                if b64:
                    # We don't know the exact mime; PNG works for most extracted images
                    md_parts.append(f"\n![page {p.index+1} image {i}](data:image/png;base64,{b64})\n")
    return "\n".join(md_parts).strip()


def response_to_docx_bytes(resp) -> bytes:
    if not HAS_PYTHON_DOCX:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    doc = DocxDocument()
    for p in resp.pages:
        doc.add_heading(f"Page {p.index + 1}", level=2)
        if getattr(p, "markdown", None):
            # DOCX doesn't support Markdown; add as plain paragraphs (simple approach)
            for line in p.markdown.splitlines():
                doc.add_paragraph(line)
        images = getattr(p, "images", []) or []
        for im in images:
            b64 = getattr(im, "image_base64", None)
            if b64:
                bio = io.BytesIO(base64.b64decode(b64))
                try:
                    doc.add_picture(bio, width=Inches(6))
                except Exception:
                    # Fallback without resizing if pillow can't infer dimensions
                    bio.seek(0)
                    doc.add_picture(bio)
        # Page break between pages
        doc.add_page_break()
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# -----------------------------
# UI
# -----------------------------
# --- Page config ---
st.set_page_config(
    page_title="PDF → Markdown / DOCX", 
    page_icon="📚", 
    layout="wide"
)

# --- Title with custom color ---
# st.markdown("""
# <style>
# .block-container {
#     padding-top: 1rem;
#     padding-bottom: 0rem;
#     padding-left: 5rem;
#     padding-right: 5rem;
# }
# </style>
# """, unsafe_allow_html=True)

# --- Title with custom color ---
st.markdown(
    "<h1 style='text-align: center; color: #2a8be9;'>PDF → Markdown / DOCX</h1>",
    unsafe_allow_html=True
)

# st.caption("Convert PDF documents into editable Markdown or Word files using Mistral OCR.")

# --- UI Layout ---
col1, col2 = st.columns([1, 1], gap="large")

# --- Left Column (Input & Controls) ---
with col1:
    # Check for API key at the top
    api_key = get_api_key()
    if not api_key:
        st.warning("Please set your Mistral API key in `st.secrets['MISTRAL_API_KEY']` or as the environment variable `MISTRAL_API_KEY`.")

    uploaded = st.file_uploader("Upload a PDF", type=["pdf"], accept_multiple_files=False)

    want_docx = st.checkbox("Also produce DOCX (Word)", value=False)

    run = st.button("Start OCR")

# --- Right Column (Output & Welcome Message) ---
with col2:
    if 'md' in st.session_state:
        # Show download buttons if conversion is done
        st.subheader("Conversion Complete!")
        st.download_button("Download Markdown (.md)", data=st.session_state.md.encode("utf-8"), file_name="ocr_output.md", mime="text/markdown")
        if st.session_state.get('docx_bytes'):
            st.download_button("Download Word (.docx)", data=st.session_state.docx_bytes, file_name="ocr_output.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # Placeholder for the main content
    main_content_placeholder = st.empty()
    
    # Show welcome message or conversion results
    if 'md' not in st.session_state:
        with main_content_placeholder:
            st.markdown("""
### 📚 PDF to Markdown & DOCX Converter

This tool **converts PDF books** those consisting of scanned images into **Markdown**, and optionally into **Word (DOCX) files**.  

All thanks to [**Mistral OCR**](https://huggingface.co/mistralai/Mistral-OCR) for providing a **highly accurate OCR model**, and to [**Streamlit**](https://streamlit.io) for making a **simple and elegant web interface** to run this tool seamlessly.  

---

### ⚡ How It Works

1. **Upload your PDF** from the left-hand panel.  
2. **Choose output format**: Markdown (default) or Word (DOCX).  
3. **Click "Start OCR"** to begin processing.  

The extracted text will appear here once processing completes, ready to use.

---

> 📊 OCR Accuracy
> - **English PDFs:** ~98% accuracy ✅  
> - **Arabic PDFs:** ~92% accuracy ✅  
> - Note: Accuracy may vary depending on image quality, font style, and PDF layout.
""")
    else:
        # Show the converted markdown
        with main_content_placeholder:
            st.subheader("Preview (Markdown)")
            st.markdown(st.session_state.md, unsafe_allow_html=True)

# --- Conversion Logic (runs after button click) ---
if run:
    # Use flags to manage flow without stopping
    is_ready_to_run = True

    if not uploaded:
        with col1:
            st.error("Please upload a PDF first.")
        is_ready_to_run = False
    
    if not api_key:
        with col1: # Placing this error in the input column is better UX
            st.error("No API key found. Add it to Streamlit Secrets or your environment.")
        is_ready_to_run = False

    if is_ready_to_run:
        pdf_bytes = uploaded.getvalue()

        # Clear previous state
        if 'md' in st.session_state:
            del st.session_state.md
        if 'docx_bytes' in st.session_state:
            del st.session_state.docx_bytes

        # Client
        client = Mistral(api_key=api_key)

        # Progress animation while the request runs in a separate thread
        progress = st.progress(0)
        status = st.empty()

        def _task():
            return run_ocr(client, pdf_bytes)

        with ThreadPoolExecutor(max_workers=1) as ex:
            fut = ex.submit(_task)
            pct = 0
            while not fut.done():
                pct = (pct + 3) % 100
                progress.progress(pct / 100)
                status.info("Processing your PDF with Mistral OCR...")
                time.sleep(0.08)
            try:
                resp = fut.result()
            except Exception as e:
                progress.empty()
                status.empty()
                st.error("An unexpected error occurred while calling Mistral OCR.")
                st.exception(e)
                # Keep running to show the footer
            
        # Finish progress
        progress.progress(1.0)
        status.success("Done! Rendering results...")

        # Compose Markdown and store it
        md = response_to_markdown(resp)
        st.session_state.md = md

        # Generate DOCX and store it if requested
        if want_docx:
            if not HAS_PYTHON_DOCX:
                st.error("`python-docx` is not installed. Run: `pip install python-docx`")
            else:
                docx_bytes = response_to_docx_bytes(resp)
                st.session_state.docx_bytes = docx_bytes
        
        # Re-run to update the UI
        st.rerun()

# --- Footer ---
st.markdown("""
<style>
.footer {
    position: fixed;
    left: 0;
    bottom: 0;
    width: 100%;
    background-color: #0e1117;
    color: #ffffff;
    text-align: center;
    padding: 12px 0;
    font-size: 16px;
    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    border-top: 1px solid #262730;
}

.footer a {
    text-decoration: none;
    margin: 0 15px;
    font-weight: 500;
    transition: opacity 0.2s ease;
}

.footer a.whatsapp {
    color: #25D366;
}

.footer a.linkedin {
    color: #2084E7;
}

.footer a:hover {
    opacity: 0.8;
    text-decoration: underline;
}
</style>

<div class="footer">
    <span>Developed by: Mohammed Babaqi</span>
    <!--
    <a href="https://wa.me/967778558924" target="_blank" class="whatsapp">WhatsApp</a>
            -->
    <a href="https://www.linkedin.com/in/mohammedbabaqi" target="_blank" class="linkedin">LinkedIn</a>
</div>

""", unsafe_allow_html=True)
